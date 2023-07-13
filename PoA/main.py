from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from tkinter import *
from tkinter import ttk
from tkinter.font import nametofont
import os
import xlrd
import openpyxl
from openpyxl import Workbook
import openpyxl

document = Document('vzor.docx')
font_styles = document.styles


font_charstyle = font_styles.add_style('parstyle', WD_STYLE_TYPE.CHARACTER)

font_object = font_charstyle.font
font_object.size = Pt(12)
font_object.name = 'Arial'

list_of_names = []
final_list = []
list_of_table = []

# #souborry XLSX

workbook = openpyxl.load_workbook("students_list.xlsx")
# Define variable to read the active sheet:
worksheet = workbook.active

#Iterate the loop to read the cell values
for i in range(1, worksheet.max_row):
    row_index = 0
    for col in worksheet.iter_cols(1, worksheet.max_column):


        if col[i].value!= None:

            one_row = col[i].value
            list_of_table.append(one_row)
            if row_index == 9:
                final_list.append(list_of_table)
                row_index = 0
                list_of_table = []
            else:
                row_index = row_index + 1



document = Document('vzor.docx')
font_styles = document.styles

style = font_styles.add_style('parstyle', WD_STYLE_TYPE.PARAGRAPH)


def create_diplom(name, course, fordate, todate, subsidy, teacher_name):
    document = Document('vzor.docx')
    for p in document.paragraphs:

        if "ftx1" in p.text:

            p.text = ""
            p.add_run(name, style='CommentsStyle').bold = True
        elif "Anglický jazyk A0.5" in p.text:
            p.text = ""
            p.add_run(course, style='CommentsStyle').bold = True
        elif "konaný v termínu" in p.text:
            font_styles = document.styles

            font_charstyle = font_styles.add_style('parstyle', WD_STYLE_TYPE.CHARACTER)

            font_object = font_charstyle.font
            font_object.size = Pt(12)
            font_object.name = 'Arial'
            p.text = ""
            p.add_run(f"	konaný v termínu od: {fordate}		do: {todate}	", style='parstyle').bold = True
        elif "Časová dotace kurzu" in p.text:
                p.text = ""
                p.add_run(f"Časová dotace kurzu: {subsidy}	 vyučovacích hodin", style='parstyle').bold = True
        # elif "délka vyučovací" in p.text:
        #     p.text = ""
        #     p.add_run(f"délka vyučovací hodiny {time} minut", style='parstyle').bold = False

        elif "Martina Vydrová" in p.text:
            p.text =""
            p.add_run(f"Lord Voldemort                         Albus Brumbál 	   {teacher_name}   	 Vedoucí kurzu                            Ředitel školy   	Lektor", style='parstyle').bold = True

# zadaná cesta do složky nutno měnit při změně cesty
    filename = f'Potvrzeni o absolvovani {name}_{course}.docx'
    # ---musíte zadat cestu k souboru!!!
    filepath = f'C:/vloz cestu k souboru/docx_create/PoA/potvrzeni/'+filename
    document.save(filepath)
    create_label.insert(END, f"vytvořen {filename}")
def create_all_diploms():
    for one_list in final_list:
        one_name = one_list[9]
        one_name = one_name.replace("\n", "")
        one_course = one_list[2]
        test_date = one_list[5]
        one_for_date = (test_date.strftime("%d")+"."+test_date.strftime("%m")+"."+test_date.strftime("%Y"))
        test_date = one_list[6]
        one_to_date = (test_date.strftime("%d")+"."+test_date.strftime("%m")+"."+test_date.strftime("%Y"))
        one_subsidy = one_list[7]*one_list[8]/60


        # one_time = one_list[8]
        one_teacher_name = one_list[1]
        create_diplom(one_name, one_course, one_for_date, one_to_date, one_subsidy, one_teacher_name)


def open_docx():
    patch = (create_label.get(ANCHOR))
    patch = patch.replace("vytvořen ", "")
    # ---musíte zadat cestu k souboru!!!
    os.startfile(f"C:/vloz cestu k souboru/docx_create/PoA/potvrzeni/{patch}")

# Grafické rozhraní

#Okna
window = Tk()
window.geometry("700x500+100+40")
window.resizable(False, False)
window.title("Docx soubory")


#Tlačítka
create_but = Button(width=20, text="Vytvoř docx", font=("Ariel", 14), command=create_all_diploms)
create_but.grid(row=0, column=0, pady=(10, 0), padx=2, sticky=W)
open_but = Button(width=20, text="Otevři docx", font=("Ariel", 14), command=open_docx)
open_but.grid(row=1, column=0, pady=(10, 0), padx=2, sticky=W)

#label pro vypsaní vytvořench souborů
create_label = Listbox(width=100, font=("Arial", 12))
create_label.grid(row=2, column=0, sticky=S+W)
#hlovní cyklus
window.mainloop()